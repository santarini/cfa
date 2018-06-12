import csv
import os
from docx import Document
from docx.shared import Inches

#create AnswerKey document
documentAnswerKey = Document()

#create BlankQuiz document
documentBlankQuiz = Document()

#add a AnswerKey title
documentAnswerKey.add_heading('CFA Formula Quiz Answer Key', 0)

#add a BlankQuiz title
documentBlankQuiz.add_heading('CFA Formula Quiz', 0)

#BEGIN LOOP
#open CSV
with open("guide.csv") as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        #find row marked Count
        NumberRef = (row['NumberRef'])
        
        #find row marked FormulaName
        FormulaName = (row['FormulaName'])
        
        #print question onto AnswerKey
        documentAnswerKey.add_paragraph(str(NumberRef) + ") " + FormulaName)

        #print question onto BlankQuiz
        documentBlankQuiz.add_paragraph(str(NumberRef) + ") " + FormulaName)

        #create an empty space on Blank Quiz
        documentBlankQuiz.add_paragraph()
        
        #add the image (aka. Answer) to AnswerKey
        documentAnswerKey.add_picture(r'C:\Users\m4k04\Desktop\workspace\imageDOc\formulaImages\/{}.png'.format(NumberRef), width=Inches(1.25))
    #next
#END LOOP

documentAnswerKey.save('AnswerKey.docx')
documentBlankQuiz.save('FomrulaQuiz.docx')
